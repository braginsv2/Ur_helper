from docx import Document
import os
from database import DatabaseManager
import psycopg2
import psycopg2.extras
from datetime import datetime
import pandas as pd
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils.dataframe import dataframe_to_rows
import num2words
from datetime import datetime, timedelta
def replace_words_in_word(search_words, replace_words, input_path, output_path):
    """
    Функция для замены и создания файла
    arg[0] = ["Список слов, которые нужно заменить"]
    arg[1] = ["Список слов, которые нужно вставить"]
    arg[2] ="Путь к шаблону"
    arg[3]= "Путь к создаваемому файлу"
    """
    
    try:
        if len(search_words) != len(replace_words):
            print("Количество слов не совпадает!")
            return False
        
        if not os.path.exists(input_path):
            print(f"Файл не найден: {input_path}")
            return False
        
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        doc = Document(input_path)
        replacements = dict(zip(search_words, replace_words))
        replacement_count = 0
        
        def replace_preserving_format(paragraph):
            nonlocal replacement_count
            
            # Собираем информацию о форматировании каждого символа
            char_formatting = []
            text_parts = []
            
            for run in paragraph.runs:
                for char in run.text:
                    char_formatting.append({
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'underline': run.font.underline,
                        'color': run.font.color.rgb
                    })
                text_parts.append(run.text)
            
            full_text = ''.join(text_parts)
            new_text = full_text
            
            # Применяем замены
            for search_word, replace_word in replacements.items():
                if search_word in new_text:
                    new_text = new_text.replace(search_word, replace_word)
                    replacement_count += 1
            
            if new_text != full_text:
                paragraph.clear()
                
                if char_formatting and len(char_formatting) >= len(full_text):
                    current_format = None
                    current_run = None
                    
                    for i, char in enumerate(new_text):

                        if i < len(char_formatting):
                            char_format = char_formatting[min(i, len(char_formatting) - 1)]
                        else:
                            char_format = char_formatting[-1] if char_formatting else {}
                        
                        if current_format != char_format:
                            current_run = paragraph.add_run()
                            if char_format.get('font_name'):
                                current_run.font.name = char_format['font_name']
                            if char_format.get('font_size'):
                                current_run.font.size = char_format['font_size']
                            if char_format.get('bold'):
                                current_run.font.bold = char_format['bold']
                            if char_format.get('italic'):
                                current_run.font.italic = char_format['italic']
                            if char_format.get('underline'):
                                current_run.font.underline = char_format['underline']
                            if char_format.get('color'):
                                current_run.font.color.rgb = char_format['color']
                            current_format = char_format
                        
                        current_run.text += char
                else:
                    paragraph.add_run(new_text)
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                replace_preserving_format(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            replace_preserving_format(paragraph)
        

        
        doc.save(output_path)
        print(f"✅ Документ сохранен: {output_path}")
        print(f"📊 Замен выполнено: {replacement_count}")
        return True
        
    except Exception as e:
        print(f"❌ Ошибка: {str(e)}")
        return False
    
def get_next_business_date():
    current_date = datetime.now().date()
    weekday = current_date.weekday()
    
    if weekday == 5:  # Суббота
        next_date = current_date + timedelta(days=2)
    elif weekday == 6:  # Воскресенье
        next_date = current_date + timedelta(days=1)
    else:  # Понедельник-пятница
        next_date = current_date
    
    # Возвращаем строку в формате ДД.ММ.ГГГГ
    return next_date.strftime('%d.%m.%Y')

# Использование
print(get_next_business_date())
        
def create_fio_data_file(data_dict):
    """
    Создает персонализированный файл данных на основе словаря с учетом пустых значений.
    """
    if 'fio' not in data_dict:
        return "Ошибка: В словаре отсутствует обязательный ключ 'fio'"
    
    fio = data_dict['fio']
    client_id = data_dict['client_id']
    # Создаем папку fio, если она не существует
    fio_dir = 'clients/'+str(client_id)
    if not os.path.exists(fio_dir):
        os.makedirs(fio_dir)
    
    # Путь к файлу
    file_path = os.path.join(fio_dir, f"{fio}_data.txt")
    
    # Если файл существует, удаляем его
    if os.path.exists(file_path):
        os.remove(file_path)
    
    # Читаем шаблон из data.txt
    template_path = 'data.txt'
    if not os.path.exists(template_path):
        return "Ошибка: Файл data.txt не найден в корневой директории"
    
    try:
        with open(template_path, 'r', encoding='utf-8') as template_file:
            lines = template_file.readlines()
        
        # Обрабатываем каждую строку
        result_lines = []
        for line in lines:
            line = line.strip()
            if ':' in line:
                # Разделяем название поля и переменную
                field_name, variable = line.split(':', 1)
                field_name = field_name.strip()
                variable = variable.strip()
                
                # Если переменная есть в словаре и значение не пустое/None
                if variable in data_dict:
                    value = data_dict[variable]
                    # Проверяем, что значение не пустое, не None и не равно 0 для числовых полей
                    if value is not None and str(value).strip() != '' and value != 0:
                        result_lines.append(f"{field_name}: {value}")
                # Иначе пропускаем эту строку (удаляем)
            else:
                # Если в строке нет двоеточия, оставляем как есть
                if line:  # Пропускаем пустые строки
                    result_lines.append(line)
        
        # Записываем результат в новый файл
        with open(file_path, 'w', encoding='utf-8') as output_file:
            for line in result_lines:
                output_file.write(line + '\n')
        
        return f"Файл успешно создан: {file_path}"
        
    except Exception as e:
        return f"Ошибка при обработке файла: {str(e)}"

def export_clients_db_to_excel(db_path='clients.db', output_path='clients_export.xlsx'):
    """
    Экспортирует данные из базы данных PostgreSQL в Excel файл
    
    Args:
        db_path (str): Не используется (оставлен для совместимости)
        output_path (str): Путь для сохранения Excel файла
    """
    
    # Словарь соответствия: русское название -> название поля в БД
    column_mapping = {
        '№ Клиента': 'client_id',
        'Статус': 'status',
        'Город': 'city',
        'Клиент ФИО': 'fio',
        'Дата ДТП': 'date_dtp',
        'Марка, модель клиента': 'marks',
        'Номер авто клиента': 'car_number',
        'Страховая компания': 'insurance',
        'Виновник ФИО Полностью': 'fio_culp',
        'Марка, модель виновника': 'marks_culp',
        'Номер авто виновника': 'number_auto_culp',
        'Дата заявления в страховую': 'date_ins',
        'Дата заявления в СТО': 'date_zayav_sto',
        'Дата Составления претензии': 'date_pret',
        'Дата составления заявления омбуцмену': 'date_ombuc',
        'Дата искового заявления': 'date_isk',
        'Суд': 'sud',
        'ID администратора': 'user_id',
        'ФИО администратора': 'admin_fio',
    }
    
    try:
        import pandas as pd
        import openpyxl
        from openpyxl.utils.dataframe import dataframe_to_rows
        from openpyxl.styles import Font, PatternFill, Alignment
        
        # Используем DatabaseManager для подключения к PostgreSQL
        db = DatabaseManager()
        
        with db.get_connection() as conn:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cursor:
                # Получаем список доступных колонок в таблице clients
                cursor.execute("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name = 'clients'
                    ORDER BY ordinal_position
                """)
                available_columns = [row['column_name'] for row in cursor.fetchall()]
                
                print(f"Доступные поля в таблице: {available_columns}")
                
                # Создаем список полей для выборки (только те, что есть в БД)
                select_columns = []
                russian_headers = []
                
                for rus_name, db_name in column_mapping.items():
                    if db_name == 'admin_fio':
                        # Специальный случай для ФИО администратора
                        select_columns.append('a.fio as admin_fio')
                        russian_headers.append(rus_name)
                    elif db_name in available_columns:
                        select_columns.append(f'c.{db_name}')
                        russian_headers.append(rus_name)
                    else:
                        print(f"Поле '{db_name}' ({rus_name}) не найдено в таблице")
                
                if not select_columns:
                    print("Не найдено ни одного совпадающего поля!")
                    return False
                
                # Выполняем запрос к базе данных с JOIN для получения данных администратора
                query = f"""
                SELECT {', '.join(select_columns)}
                FROM clients c
                LEFT JOIN admins a ON c.user_id = a.user_id::text AND a.is_active = true
                ORDER BY c.created_at DESC
                """
                
                cursor.execute(query)
                results = cursor.fetchall()
                
                if not results:
                    print("Нет данных для экспорта!")
                    return False
                
                # Преобразуем в DataFrame
                df = pd.DataFrame([dict(row) for row in results])
                
                # Переименовываем колонки на русские названия
                df.columns = russian_headers
                
                print(f"Загружено {len(df)} записей с {len(df.columns)} полями")
                
                # Создаем Excel файл с форматированием
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Клиенты"
                
                # Добавляем данные в лист
                for r in dataframe_to_rows(df, index=False, header=True):
                    ws.append(r)
                
                # Форматирование заголовков
                header_font = Font(bold=True, color="FFFFFF")
                header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                
                for col in range(1, len(russian_headers) + 1):
                    cell = ws.cell(row=1, column=col)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                
                # Автоматическая ширина колонок
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    
                    adjusted_width = min(max_length + 2, 50)  # Максимальная ширина 50
                    ws.column_dimensions[column_letter].width = adjusted_width
                
                # Замораживаем первую строку
                ws.freeze_panes = "A2"
                
                # Сохраняем файл
                wb.save(output_path)
                
                print(f"Экспорт завершен успешно!")
                print(f"Файл сохранен: {output_path}")
                print(f"Экспортировано записей: {len(df)}")
                print(f"Количество полей: {len(df.columns)}")
                
                return True
        
    except Exception as e:
        print(f"Ошибка при экспорте: {e}")
        import traceback
        traceback.print_exc()
        return False
def load_field_mapping_from_data_file():
    """
    Загружает маппинг полей из файла data.txt
    Возвращает словарь: русское_название -> название_поля_в_бд
    """
    field_mapping = {}
    
    try:
        with open('data.txt', 'r', encoding='utf-8') as file:
            lines = file.readlines()
        
        for line in lines:
            line = line.strip()
            if ':' in line:
                # Разделяем название поля и переменную
                field_name, variable = line.split(':', 1)
                field_name = field_name.strip().lower()
                variable = variable.strip()
                
                # Добавляем в маппинг
                field_mapping[field_name] = variable
        
        print(f"Загружено {len(field_mapping)} полей из data.txt")
        return field_mapping
        
    except Exception as e:
        print(f"Ошибка при загрузке маппинга полей: {e}")
        # Возвращаем базовый маппинг на случай ошибки
        return {
            'паспорт серия клиента': 'seria_pasport',
            'паспорт номер клиента': 'number_pasport',
            'паспорт выдан клиента': 'where_pasport',
            'паспорт когда выдан клиента': 'when_pasport',
            'дата рождения клиента': 'date_of_birth',
            'город клиента': 'city',
            'адрес клиента': 'address',
            'индекс клиента': 'index_postal',
            'телефон клиента': 'number',
            'марка модель клиента': 'marks',
            'номер авто клиента': 'car_number',
            'год авто клиента': 'year_auto',
            'страховая компания': 'insurance',
            'серия полиса': 'seria_insurance',
            'номер полиса': 'number_insurance',
            'дата полиса': 'date_insurance',
            'дата дтп': 'date_dtp',
            'время дтп': 'time_dtp',
            'адрес дтп': 'address_dtp',
            'фио виновника': 'fio_culp',
            'марка модель виновника': 'marks_culp',
            'номер авто виновника': 'number_auto_culp'
        }
def edit_files(files, data):
    for i in files:
        if "1. Обложка дела.docx" == i:
            try:
                replace_words_in_word(["{{ Дата_ДТП }}", "{{ Время_ДТП }}", "{{ Адрес_ДТП }}", 
                                "{{ Марка_модель }}", "{{ Nавто_клиента }}", "{{ Год }}","{{ NКлиента }}", "{{ ФИО }}",
                                "{{ Страховая }}", "{{ винФИО }}"],
                                [str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]), str(data["marks"]), str(data["car_number"]),
                                    str(data['year']),str(data['client_id']), str(data["fio"]), str(data["insurance"]), str(data["fio_culp"])],
                                    "Шаблоны/1. ДТП/1. На ремонт/1. Обложка дела.docx",
                                    "clients/"+str(data["client_id"])+"/Документы/"+"1. Обложка дела.docx")
            except Exception as e:
                print(e)
                print(i)
                print("1. Обложка дела.docx")
        elif "2. Юр договор.docx" == i:
            try:
                replace_words_in_word(
                ["{{ Год }}", "{{ NКлиента }}", "{{ Город }}", "{{ Дата }}", "{{ ФИО }}", 
                 "{{ ДР }}", "{{ Паспорт_серия }}", "{{ Паспорт_номер }}", "{{ Паспорт_выдан }}", 
                 "{{ Паспорт_когда }}", "{{ Индекс }}", "{{ Адрес }}", "{{ Дата_ДТП }}", 
                 "{{ Время_ДТП }}", "{{ Адрес_ДТП }}", "{{ ФИОк }}"],
                [str(data['year']), str(data["client_id"]), str(data["city"]), 
                 str(datetime.now().strftime("%d.%m.%Y")), str(data["fio"]), 
                 str(data["date_of_birth"]), str(data["seria_pasport"]), 
                 str(data["number_pasport"]), str(data["where_pasport"]),
                 str(data["when_pasport"]), str(data["index_postal"]), 
                 str(data["address"]), str(data["date_dtp"]), 
                 str(data["time_dtp"]), str(data["address_dtp"]), 
                 str(data['fio_k'])],
                "Шаблоны/1. ДТП/1. На ремонт/2. Юр договор.docx",
                f"clients/"+str(data["client_id"])+"/Документы/2. Юр договор.docx"
            )
            except Exception as e:
                print(e)
                print(i)
                print("2. Юр договор.docx")
        elif "3a Заявление в Страховую ФЛ собственник с эвакуатором.docx" == i:
            try:
                replace_words_in_word(
                ["{{ Страховая }}", "{{ ФИО }}", "{{ Паспорт_серия }}", 
                "{{ Паспорт_номер }}", "{{ ДР }}", "{{ Индекс }}",
                "{{ Адрес }}", "{{ Марка_модель }}", "{{ Год_авто }}", "{{ Nавто_клиента }}", "{{ Документ }}",
                "{{ Док_серия }}", "{{ Док_номер }}", "{{ Док_когда }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                "{{ Адрес_ДТП }}", "{{ винФИО }}", "{{ Марка_модель_виновника }}", "{{ Серия_полиса }}",
                "{{ Номер_полиса }}", "{{ Город }}", "{{ Дата }}", "{{ Место }}"],
                [str(data["insurance"]), str(data["fio"]), str(data["seria_pasport"]), str(data["number_pasport"]),
                str(data["date_of_birth"]), str(data["index_postal"]), str(data["address"]),
                str(data["marks"]), str(data["year_auto"]), str(data["car_number"]), str(data["docs"]), 
                str(data["seria_docs"]), str(data["number_docs"]), str(data["data_docs"]), 
                str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]),
                str(data["fio_culp"]), str(data["marks_culp"]), str(data["seria_insurance"]),
                str(data["number_insurance"]), str(data["city"]), str(datetime.now().strftime("%d.%m.%Y")), str(data["city_birth"])],
                "Шаблоны/1. ДТП/1. На ремонт/3. Заявление в страховую после ДТП/3a Заявление в Страховую ФЛ собственник с эвакуатором.docx",
                f"clients/{data['client_id']}/Документы/3a Заявление в Страховую ФЛ собственник с эвакуатором.docx")
            except Exception as e:
                print(e)
                print(i)
                print("3a Заявление в Страховую ФЛ собственник с эвакуатором.docx")
        elif "3b Заявление в Страховую ФЛ собственник без эвакуатора.docx" == i:
            try:
                replace_words_in_word(
                ["{{ Страховая }}", "{{ ФИО }}", "{{ Паспорт_серия }}", 
                "{{ Паспорт_номер }}", "{{ ДР }}", "{{ Индекс }}",
                "{{ Адрес }}", "{{ Марка_модель }}", "{{ Год_авто }}", "{{ Nавто_клиента }}", "{{ Документ }}",
                "{{ Док_серия }}", "{{ Док_номер }}", "{{ Док_когда }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                "{{ Адрес_ДТП }}", "{{ винФИО }}", "{{ Марка_модель_виновника }}", "{{ Серия_полиса }}",
                "{{ Номер_полиса }}", "{{ Город }}", "{{ Дата }}", "{{ Место }}"],
                [str(data["insurance"]), str(data["fio"]), str(data["seria_pasport"]), str(data["number_pasport"]),
                str(data["date_of_birth"]), str(data["index_postal"]), str(data["address"]),
                str(data["marks"]), str(data["year_auto"]), str(data["car_number"]), str(data["docs"]), 
                str(data["seria_docs"]), str(data["number_docs"]), str(data["data_docs"]), 
                str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]),
                str(data["fio_culp"]), str(data["marks_culp"]), str(data["seria_insurance"]),
                str(data["number_insurance"]), str(data["city"]), str(datetime.now().strftime("%d.%m.%Y")), str(data["city_birth"])],
                "Шаблоны/1. ДТП/1. На ремонт/3. Заявление в страховую после ДТП/3b Заявление в Страховую ФЛ собственник без эвакуатора.docx",
                f"clients/{data['client_id']}/Документы/3b Заявление в Страховую ФЛ собственник без эвакуатора.docx")
            except Exception as e:
                print(e)
                print(i)
                print("3b Заявление в Страховую ФЛ собственник без эвакуатора.docx")
        elif "4. Заявление о проведении дополнительного осмотра автомобиля представитель.docx" == i:
            try:
                replace_words_in_word(["{{ Страховая }}", "{{ Город }}", "{{ ФИО }}", 
                                "{{ ДР }}", "{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                                "{{ Паспорт_когда }}", 
                                "{{ NДоверенности }} ", "{{ Дата_доверенности }}", "{{ Представитель }}", "{{ Телефон_представителя }}",
                                "{{ Nакта_осмотра }}", "{{ Дата }}","{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                                "{{ Адрес_ДТП }}", "{{ Дата_осмотра }}",
                                "{{ Дата_свое_СТО }}","{{ Время_свое_СТО }}","{{ Адрес_свое_СТО }}", "{{ Телефон }}", 
                                "{{ Дата_заявления_доп_осмотр }}"],
                                [str(data["insurance"]), str(data["city"]), str(data["fio"]), str(data["date_of_birth"]),
                                    str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                                    str(data["N_dov_not"]), str(data["data_dov_not"]), str(data["fio_not"]), str(data["number_not"]),
                                    str(data["Na_ins"]), str(data["date_ins"]), str(data["date_dtp"]), str(data["time_dtp"]),str(data["address_dtp"]), 
                                    str(data["date_Na_ins"]), str(data["date_sto_main"]),
                                    str(data["time_sto_main"]), str(data["address_sto_main"]), str(data["number"]),
                                    str(data["data_dop_osm"])],
                                    "Шаблоны/1. ДТП/1. На ремонт/4. Заявление о проведении доп осмотра/4. Заявление о проведении дополнительного осмотра автомобиля представитель.docx",
                                    "clients/"+str(data["client_id"])+"/Документы/"+"4. Заявление о проведении дополнительного осмотра автомобиля представитель.docx")
            except Exception as e:
                print(e)
                print(i)
                print("4. Заявление о проведении дополнительного осмотра автомобиля представитель.docx")
        elif "4. Заявление о проведении дополнительного осмотра автомобиля.docx" == i:
            try:
                replace_words_in_word(["{{ Страховая }}", "{{ Город }}", "{{ ФИО }}", 
                                "{{ ДР }}", "{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                                "{{ Паспорт_когда }}", "{{ Nакта_осмотра }}", "{{ Дата }}","{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                                "{{ Адрес_ДТП }}", "{{ Дата_осмотра }}",
                                "{{ Дата_свое_СТО }}","{{ Время_свое_СТО }}","{{ Адрес_свое_СТО }}", "{{ Телефон }}",
                                "{{ Дата_заявления_доп_осмотр }}"],
                                [str(data["insurance"]), str(data["city"]), str(data["fio"]), str(data["date_of_birth"]),
                                    str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                                    str(data["Na_ins"]), str(data["date_ins"]), str(data["date_dtp"]), str(data["time_dtp"]),str(data["address_dtp"]), 
                                    str(data["date_Na_ins"]), str(data["date_sto_main"]),
                                    str(data["time_sto_main"]), str(data["address_sto_main"]), str(data["number"]),
                                    str(data["data_dop_osm"])],
                                    "Шаблоны/1. ДТП/1. На ремонт/4. Заявление о проведении доп осмотра/4. Заявление о проведении дополнительного осмотра автомобиля.docx",
                                    "clients/"+str(data["client_id"])+"/Документы/"+"4. Заявление о проведении дополнительного осмотра автомобиля.docx")
            except Exception as e:
                print(e)
                print(i)
                print("4. Заявление о проведении дополнительного осмотра автомобиля.docx")
        elif "5. Запрос в страховую о выдаче акта и расчёта представитель.docx" == i:
            try:
                replace_words_in_word(["{{ Страховая }}", "{{ Город }}", "{{ ФИО }}", 
                            "{{ ДР }}", "{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}", "{{ Паспорт_когда }}",
                            "{{ NДоверенности }}","{{ Дата_доверенности }}", "{{ Представитель }}","{{ Телефон_представителя }}",
                            "{{ Дата_ДТП }}","{{ Время_ДТП }}", "{{ Адрес_ДТП }}","{{ Марка_модель }}","{{ Nавто_клиента }}",
                            "{{ Марка_модель_виновника }}", "{{ Nавто_виновник }}", "{{ Телефон }}"],
                            [str(data['insurance']), str(data["city"]), str(data["fio"]), str(data["date_of_birth"]), str(data["seria_pasport"]),
                            str(data["number_pasport"]), str(data["where_pasport"]),
                            str(data["when_pasport"]),str(data["N_dov_not"]), str(data["data_dov_not"]),str(data["fio_not"]),str(data["number_not"]),
                            str(data["date_dtp"]), str(data["time_dtp"]), 
                            str(data["address_dtp"]), str(data["marks"]), str(data["car_number"]), 
                            str(data["marks_culp"]), str(data["number_auto_culp"]), str(data["number"])],
                            "Шаблоны/1. ДТП/1. На ремонт/5. Запрос в страховую о выдаче акта и расчета/5. Запрос в страховую о выдаче акта и расчёта представитель.docx",
                                "clients/"+str(data["client_id"])+"/Документы/"+"5. Запрос в страховую о выдаче акта и расчёта представитель.docx")
            except Exception as e:
                print(e)
                print(i)
                print("5. Запрос в страховую о выдаче акта и расчёта представитель.docx")
        elif "5. Запрос в страховую о выдаче акта и расчёта.docx" == i:
            try:
                replace_words_in_word(["{{ Страховая }}", "{{ Город }}", "{{ ФИО }}", 
                            "{{ ДР }}", "{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}", "{{ Паспорт_когда }}",
                            "{{ Дата_ДТП }}","{{ Время_ДТП }}", "{{ Адрес_ДТП }}","{{ Марка_модель }}","{{ Nавто_клиента }}",
                            "{{ Марка_модель_виновника }}", "{{ Nавто_виновник }}","{{ ФИОк }}", "{{ Телефон }}"],
                            [str(data['insurance']), str(data["city"]), str(data["fio"]), str(data["date_of_birth"]), str(data["seria_pasport"]),
                            str(data["number_pasport"]), str(data["where_pasport"]),
                            str(data["when_pasport"]), str(data["date_dtp"]), str(data["time_dtp"]), 
                            str(data["address_dtp"]), str(data["marks"]), str(data["car_number"]), 
                            str(data["marks_culp"]), str(data["number_auto_culp"]), str(data["fio_k"]), str(data["number"])],
                            "Шаблоны/1. ДТП/1. На ремонт/5. Запрос в страховую о выдаче акта и расчета/5. Запрос в страховую о выдаче акта и расчёта.docx",
                                "clients/"+str(data["client_id"])+"/Документы/"+"5. Запрос в страховую о выдаче акта и расчёта.docx")
            except Exception as e:
                print(e)
                print(i)
                print("5. Запрос в страховую о выдаче акта и расчёта.docx")
        elif "6. Заявление в СТО представитель.docx" == i:
            try:
                replace_words_in_word(["{{ СТО }}", "{{ ИНН_СТО }}", "{{ Индекс_СТО }}", 
                                "{{ Адрес_СТО }}", "{{ ФИО }}","{{ ДР }}", "{{ Паспорт_серия }}",
                                "{{ Паспорт_номер }}", "{{ Паспорт_выдан }}", "{{ Паспорт_когда }}",
                                "{{ NДоверенности }}", "{{ Дата_доверенности }}", "{{ Представитель }}","{{ Телефон_представителя }}",
                                "{{ Номер_направления_СТО }}",
                                "{{ Страховая }}","{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                                "{{ Адрес_ДТП }}", "{{ Дата_предоставления_ТС }}", "{{ Марка_модель }}", "{{ Nавто_клиента }}",
                                "{{ Дата_Заявления_СТО }}", "{{ ФИОк }}", "{{ Дата }}", "{{ Телефон }}"],
                                [str(data["name_sto"]), str(data["inn_sto"]), str(data["index_sto"]),
                                    str(data["address_sto"]), str(data["fio"]),str(data["date_of_birth"]), str(data["seria_pasport"]),
                                    str(data["number_pasport"]), str(data["where_pasport"]), str(data["when_pasport"]),
                                    str(data["N_dov_not"]), str(data["data_dov_not"]),str(data["fio_not"]), str(data["number_not"]),
                                    str(data["N_sto"]),
                                    str(data["insurance"]), str(data["date_dtp"]), str(data["time_dtp"]),str(data["address_dtp"]), 
                                    str(data["date_sto"]), str(data["marks"]), str(data["car_number"]), str(data["date_zayav_sto"]),str(data["fio_k"]),
                                    str(data["date_ins"]), str(data["number"])],
                                    "Шаблоны/1. ДТП/1. На ремонт/Ремонт не произведен СТО отказала/6. Заявление в СТО представитель.docx",
                                    "clients/"+str(data["client_id"])+"/Документы/"+"6. Заявление в СТО представитель.docx")
            except Exception as e:
                print(e)
                print(i)
                print("6. Заявление в СТО представитель.docx")
        elif "6. Заявление в СТО.docx" == i:
            try:
                replace_words_in_word(["{{ СТО }}", "{{ ИНН_СТО }}", "{{ Индекс_СТО }}", 
                                "{{ Адрес_СТО }}", "{{ ФИО }}","{{ ДР }}", "{{ Паспорт_серия }}",
                                "{{ Паспорт_номер }}", "{{ Паспорт_выдан }}", "{{ Паспорт_когда }}","{{ Номер_направления_СТО }}",
                                "{{ Страховая }}","{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                                "{{ Адрес_ДТП }}", "{{ Дата_предоставления_ТС }}", "{{ Марка_модель }}", "{{ Nавто_клиента }}",
                                "{{ Дата_Заявления_СТО }}", "{{ ФИОк }}", "{{ Дата }}", "{{ Телефон }}"],
                                [str(data["name_sto"]), str(data["inn_sto"]), str(data["index_sto"]),
                                    str(data["address_sto"]), str(data["fio"]),str(data["date_of_birth"]), str(data["seria_pasport"]),
                                    str(data["number_pasport"]), str(data["where_pasport"]), str(data["when_pasport"]), str(data["N_sto"]),
                                    str(data["insurance"]), str(data["date_dtp"]), str(data["time_dtp"]),str(data["address_dtp"]), 
                                    str(data["date_sto"]), str(data["marks"]), str(data["car_number"]), str(data["date_zayav_sto"]),str(data["fio_k"]),
                                    str(data["date_ins"]), str(data["number"])],
                                    "Шаблоны/1. ДТП/1. На ремонт/Ремонт не произведен СТО отказала/6. Заявление в СТО.docx",
                                    "clients/"+str(data["client_id"])+"/Документы/"+"6. Заявление в СТО.docx")
            except Exception as e:
                print(e)
                print(i)
                print("6. Заявление в СТО.docx")
        elif "7. Заявление фин. омбудсмену при выплате без согласования.docx" == i:
            try:
                replace_words_in_word(["{{ Дата_обуцмен }}", "{{ Страховая }}","{{ Город }}", "{{ ФИО }}", 
                            "{{ ДР }}", "{{ Место }}","{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                            "{{ Паспорт_когда }}", "{{ Адрес }}", "{{ Телефон }}","{{ Серия_полиса }}","{{ Номер_полиса }}",
                            "{{ Дата_полиса }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                            "{{ Адрес_ДТП }}", "{{ Марка_модель }}", "{{ Nавто_клиента }}", "{{ Дата }}",
                            "{{ Организация }}", "{{ Nэкспертизы }}","{{ Дата_экспертизы }}", "{{ Без_учета_износа }}",
                            "{{ С_учетом_износа }}", "{{ Дата_претензии }}", "{{ Дата_ответа_на_претензию }}", "{{ Выплата_ОСАГО }}", "{{ ФИОк }}", "{{ Nв_страховую }}"],
                            [str(data["date_ombuc"]), str(data["insurance"]),str(data["city"]), str(data["fio"]), str(data["date_of_birth"]), str(data["city_birth"]),
                                str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                                str(data["address"]), str(data["number"]), str(data["seria_insurance"]), str(data["number_insurance"]),str(data["date_insurance"]), 
                                str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]), str(data["marks"]), str(data["car_number"]),
                                str(data["date_ins_pod"]), str(data["org_exp"]), str(data["Na_ins"]),str(data["date_exp"]),
                                str(data["coin_exp"]), str(data["coin_exp_izn"]),str(data["date_pret"]),
                                str(data["data_pret_otv"]), str(data["coin_osago"]),str(data["fio_k"]), str(data["Nv_ins"])],
                                "Шаблоны/1. ДТП/1. На ремонт/Выплата без согласования/7. Заявление фин. омбудсмену при выплате без согласования.docx",
                                "clients/"+str(data["client_id"])+"/Документы/"+"7. Заявление фин. омбудсмену при выплате без согласования.docx")
            except Exception as e:
                print(e)
                print(i)
                print("7. Заявление фин. омбудсмену при выплате без согласования.docx")
        elif "8. Заявление фин. омбуцмену СТО отказала.docx" == i:
            try:
                replace_words_in_word(["{{ Дата_обуцмен }}", "{{ Страховая }}", "{{ ФИО }}", 
                        "{{ ДР }}", "{{ Место }}","{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                        "{{ Паспорт_когда }}", "{{ Адрес }}", "{{ Телефон }}","{{ Серия_полиса }}","{{ Номер_полиса }}",
                        "{{ Дата_полиса }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                        "{{ Адрес_ДТП }}", "{{ Марка_модель }}", "{{ Nавто_клиента }}", "{{ Дата }}",
                        "{{ Nв_страховую }}", "{{ Дата_направления_ремонт }}","{{ Номер_направления_СТО }}", "{{ СТО }}",
                        "{{ Индекс_СТО }}", "{{ Адрес_СТО }}", "{{ Дата_предоставления_ТС }}", "{{ Дата_принятия_претензии }}", "{{ Nпринятой_претензии }}",
                        "{{ Дата_претензии }}", "{{ Банк_получателя }}", "{{ Счет_получателя }}","{{ Кор_счет_получателя }}", "{{ БИК_Банка }}", "{{ ИНН_Банка }}",
                        "{{ ФИОк }}","{{ Организация }}", "{{ Nэкспертизы }}", "{{ Дата_экспертизы }}", "{{ Без_учета_износа }}", "{{ С_учетом_износа }}", "{{ Город }}" ],
                        [str(data["date_ombuc"]), str(data["insurance"]), str(data["fio"]), str(data["date_of_birth"]), str(data["city_birth"]),
                            str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                            str(data["address"]), str(data["number"]), str(data["seria_insurance"]), str(data["number_insurance"]),str(data["date_insurance"]), 
                            str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]), str(data["marks"]), str(data["car_number"]),
                            str(data["date_ins_pod"]), str(data["Nv_ins"]), str(data["date_napr_sto"]),str(data["N_sto"]),
                            str(data["name_sto"]), str(data["index_sto"]),str(data["address_sto"]), str(data["date_sto"]),
                            str(data["data_pret_prin"]),str(data["N_pret_prin"]),str(data["date_pret"]),str(data["bank"]),str(data["bank_account"]),
                            str(data["bank_account_corr"]),str(data["BIK"]),str(data["INN"]),str(data["fio_k"]), str(data["org_exp"]),str(data["Na_ins"]),
                            str(data["date_exp"]), str(data["coin_exp"]), str(data["coin_exp_izn"]), str(data["city"])],
                            "Шаблоны/1. ДТП/1. На ремонт/Ремонт не произведен СТО отказала/8. Заявление фин. омбуцмену СТО отказала.docx",
                            "clients/"+str(data["client_id"])+"/Документы/"+"8. Заявление фин. омбуцмену СТО отказала.docx")
            except Exception as e:
                print(e)
                print(i)
                print("8. Заявление фин. омбуцмену СТО отказала.docx")
        elif "7. Заявление фин. омбудсмену СТО свыше 50 км.docx" == i:
            try:
                replace_words_in_word(["{{ Дата_обуцмен }}", "{{ Страховая }}", "{{ ФИО }}", 
                        "{{ ДР }}", "{{ Место }}","{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                        "{{ Паспорт_когда }}", "{{ Адрес }}", "{{ Телефон }}","{{ Серия_полиса }}","{{ Номер_полиса }}",
                        "{{ Дата_полиса }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                        "{{ Адрес_ДТП }}", "{{ Марка_модель }}", "{{ Nавто_клиента }}", "{{ Дата }}",
                        "{{ Nв_страховую }}", "{{ Дата_направления_ремонт }}","{{ Номер_направления_СТО }}", "{{ СТО }}",
                        "{{ Индекс_СТО }}", "{{ Адрес_СТО }}", "{{ Дата_предоставления_ТС }}", "{{ Дата_принятия_претензии }}", "{{ Nпринятой_претензии }}",
                        "{{ Дата_претензии }}", "{{ Банк_получателя }}", "{{ Счет_получателя }}","{{ Кор_счет_получателя }}", "{{ БИК_Банка }}", "{{ ИНН_Банка }}",
                        "{{ ФИОк }}","{{ Организация }}", "{{ Nэкспертизы }}", "{{ Дата_экспертизы }}", "{{ Без_учета_износа }}", "{{ С_учетом_износа }}",
                        "{{ Город }}","{{ Город_СТО }}"],
                        [str(data["date_ombuc"]), str(data["insurance"]), str(data["fio"]), str(data["date_of_birth"]), str(data["city_birth"]),
                            str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                            str(data["address"]), str(data["number"]), str(data["seria_insurance"]), str(data["number_insurance"]),str(data["date_insurance"]), 
                            str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]), str(data["marks"]), str(data["car_number"]),
                            str(data["date_ins_pod"]), str(data["Nv_ins"]), str(data["date_napr_sto"]),str(data["N_sto"]),
                            str(data["name_sto"]), str(data["index_sto"]),str(data["address_sto"]), str(data["date_sto"]),
                            str(data["data_pret_prin"]),str(data["N_pret_prin"]),str(data["date_pret"]),str(data["bank"]),str(data["bank_account"]),
                            str(data["bank_account_corr"]),str(data["BIK"]),str(data["INN"]),str(data["fio_k"]), str(data["org_exp"]),str(data["Na_ins"]),
                            str(data["date_exp"]), str(data["coin_exp"]), str(data["coin_exp_izn"]), str(data["city"]), str(data["city_sto"])],
                            "Шаблоны/1. ДТП/1. На ремонт/Ремонт не произведен СТО свыше 50км/7. Заявление фин. омбудсмену СТО свыше 50 км.docx",
                            "clients/"+str(data["client_id"])+"/Документы/"+"7. Заявление фин. омбудсмену СТО свыше 50 км.docx")
            except Exception as e:
                print(e)
                print(i)
                print("7. Заявление фин. омбудсмену СТО свыше 50 км.docx")
        elif "6. Претензия в страховую Выплата без согласования.docx" == i:
            try:
                replace_words_in_word(["{{ Страховая }}", "{{ Город }}", "{{ ФИО }}", 
                                            "{{ ДР }}", "{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                                            "{{ Паспорт_когда }}", "{{ NДоверенности }}", "{{ Дата_доверенности }}","{{ Представитель }}","{{ Телефон_представителя }}",
                                            "{{ Nакта_осмотра }}", "{{ Дата }}", "{{ Nв_страховую }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                                            "{{ Адрес_ДТП }}", "{{ Организация }}", "{{ Дата_экспертизы }}", "{{ Без_учета_износа }}",
                                            "{{ С_учетом_износа }}", "{{ Выплата_ОСАГО }}","{{ Дата_претензии }}"],
                                            [str(data["insurance"]), str(data["city"]), str(data["fio"]), str(data["date_of_birth"]),
                                                str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                                                str(data["N_dov_not"]), str(data["data_dov_not"]), str(data["fio_not"]), str(data["number_not"]),str(data["Na_ins"]), 
                                                str(data["date_ins"]), str(data["Nv_ins"]), str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]),
                                                str(data["org_exp"]), str(data["date_exp"]), str(data["coin_exp"]),str(data["coin_exp_izn"]),
                                                str(data["coin_osago"]), str(datetime.now().strftime("%d.%m.%Y"))],
                                                "Шаблоны/1. ДТП/1. На ремонт/Выплата без согласования/6. Претензия в страховую Выплата без согласования.docx",
                                                "clients/"+str(data["client_id"])+"/Документы/"+"6. Претензия в страховую Выплата без согласования.docx")
            except Exception as e:
                print(e)
                print(i)
                print("6. Претензия в страховую Выплата без согласования.docx")
        elif "7. Претензия в страховую СТО отказала.docx" == i:
            try:
                replace_words_in_word(["{{ Страховая }}", "{{ Город }}", "{{ ФИО }}", 
                                            "{{ ДР }}", "{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                                            "{{ Паспорт_когда }}", "{{ NДоверенности }}", "{{ Дата_доверенности }}","{{ Представитель }}","{{ Телефон_представителя }}",
                                            "{{ Nакта_осмотра }}", "{{ Дата }}", "{{ Nв_страховую }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                                            "{{ Адрес_ДТП }}", "{{ Дата_направления_ремонт }}", "{{ Номер_направления_СТО }}", "{{ Дата_предоставления_ТС }}",
                                            "{{ СТО }}", "{{ Дата_отказа_СТО }}","{{ Дата_претензии }}","{{ Город_СТО }}","{{ Марка_модель }}", "{{ Nавто_клиента }}"],
                                            [str(data["insurance"]), str(data["city"]), str(data["fio"]), str(data["date_of_birth"]),
                                                str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                                                str(data["N_dov_not"]), str(data["data_dov_not"]), str(data["fio_not"]), str(data["number_not"]),str(data["Na_ins"]), 
                                                str(data["date_ins"]), str(data["Nv_ins"]), str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]),
                                                str(data["date_napr_sto"]), str(data["N_sto"]), str(data["date_sto"]),str(data["name_sto"]),
                                                str(data["data_otkaz_sto"]), str(data["date_pret"]), str(data["city"]), str(data["marks"]),str(data["car_number"])],
                                                "Шаблоны/1. ДТП/1. На ремонт/Ремонт не произведен СТО отказала/7. Претензия в страховую СТО отказала.docx",
                                                "clients/"+str(data["client_id"])+"/Документы/"+"7. Претензия в страховую СТО отказала.docx")
            except Exception as e:
                print(e)
                print(i)
                print("7. Претензия в страховую СТО отказала.docx")
        elif "6. Претензия в страховую  СТО свыше 50 км.docx" == i:
            try:
                replace_words_in_word(["{{ Страховая }}", "{{ Город }}", "{{ ФИО }}", 
                                            "{{ ДР }}", "{{ Паспорт_серия }}","{{ Паспорт_номер }}", "{{ Паспорт_выдан }}",
                                            "{{ Паспорт_когда }}", "{{ NДоверенности }}", "{{ Дата_доверенности }}","{{ Представитель }}","{{ Телефон_представителя }}",
                                            "{{ Nакта_осмотра }}", "{{ Дата }}", "{{ Nв_страховую }}", "{{ Дата_ДТП }}", "{{ Время_ДТП }}", 
                                            "{{ Адрес_ДТП }}", "{{ Дата_направления_ремонт }}", "{{ Номер_направления_СТО }}",
                                            "{{ СТО }}", "{{ Индекс_СТО }}","{{ Адрес_СТО }}","{{ Город_СТО }}","{{ Номер_направления_на_ремонт }}","{{ Дата_направления }}",
                                            "{{ Марка_модель }}", "{{ Nавто_клиента }}","{{ Дата_претензии }}"],
                                            [str(data["insurance"]), str(data["city"]), str(data["fio"]), str(data["date_of_birth"]),
                                                str(data["seria_pasport"]), str(data["number_pasport"]),str(data["where_pasport"]), str(data["when_pasport"]),
                                                str(data["N_dov_not"]), str(data["data_dov_not"]), str(data["fio_not"]), str(data["number_not"]),str(data["Na_ins"]), 
                                                str(data["date_ins"]), str(data["Nv_ins"]), str(data["date_dtp"]), str(data["time_dtp"]), str(data["address_dtp"]),
                                                str(data["date_napr_sto"]), str(data["N_sto"]), str(data["name_sto"]),str(data["index_sto"]),str(data["address_sto"]),
                                                str(data["city_sto"]), str(data["N_sto"]), str(data["date_napr_sto"]), str(data["marks"]),str(data["car_number"]), str(data["date_pret"])],
                                                "Шаблоны/1. ДТП/1. На ремонт/Ремонт не произведен СТО свыше 50км/6. Претензия в страховую  СТО свыше 50 км.docx",
                                                "clients/"+str(data["client_id"])+"/Документы/"+"6. Претензия в страховую  СТО свыше 50 км.docx")
            except Exception as e:
                print(e)
                print(i)
                print("6. Претензия в страховую  СТО свыше 50 км.docx")
        else:
            pass